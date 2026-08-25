import asyncio
import hashlib
import html
import io
import json
import os
import re
import sqlite3
import smtplib
import urllib.parse
from contextlib import asynccontextmanager
from datetime import datetime, timedelta, timezone
from pathlib import Path
from typing import Optional

import requests
from docx import Document
from email.message import EmailMessage
from fastapi import FastAPI, HTTPException, Response
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

DB_PATH = os.getenv("DATABASE_PATH", "/data/job-watch.sqlite")
RESUME_BUILDER_DB = os.getenv("RESUME_BUILDER_DB", "/resume-builder-data/resume-builder.sqlite")
ANTHROPIC_MODEL = os.getenv("ANTHROPIC_MODEL", "claude-sonnet-4-20250514")
ANTHROPIC_API_KEY = os.getenv("ANTHROPIC_API_KEY") or os.getenv("ANTHROPIC_TOKEN") or ""
TELEGRAM_BOT_TOKEN = os.getenv("TELEGRAM_BOT_TOKEN", "")
TELEGRAM_CHAT_ID = os.getenv("TELEGRAM_HOME_CHANNEL") or os.getenv("TELEGRAM_CHAT_ID", "")
INDEED_MCP_URL = os.getenv("INDEED_MCP_URL", "https://mcp.indeed.com/claude/mcp")
SCHEDULER_ENABLED = os.getenv("SCHEDULER_ENABLED", "1") == "1"
ALERT_BATCH_LIMIT = int(os.getenv("ALERT_BATCH_LIMIT", "10"))
NOTIFICATION_WINDOW_HOURS = int(os.getenv("NOTIFICATION_WINDOW_HOURS", "24"))
SMTP_ENABLED = os.getenv("SMTP_ENABLED", "0") == "1"
SMTP_SERVER = os.getenv("SMTP_SERVER", "")
SMTP_PORT = int(os.getenv("SMTP_PORT", "587"))
SMTP_USERNAME = os.getenv("SMTP_USERNAME", "")
SMTP_PASSWORD = os.getenv("SMTP_PASSWORD", "")
SMTP_TO = os.getenv("SMTP_TO", "")
QUIET_HOURS_START = os.getenv("QUIET_HOURS_START", "")
QUIET_HOURS_END = os.getenv("QUIET_HOURS_END", "")
DASHBOARD_URL = os.getenv("DASHBOARD_URL", "http://10.10.10.237:8085").rstrip("/")
JSEARCH_RAPIDAPI_KEY = os.getenv("JSEARCH_RAPIDAPI_KEY") or os.getenv("RAPIDAPI_KEY", "")
JSEARCH_ENABLED = os.getenv("JSEARCH_ENABLED", "1") == "1"
JSEARCH_MONTHLY_LIMIT = int(os.getenv("JSEARCH_MONTHLY_LIMIT", "200"))
JSEARCH_PER_SCRAPE_LIMIT = int(os.getenv("JSEARCH_PER_SCRAPE_LIMIT", "1"))
JSEARCH_RAPIDAPI_HOST = os.getenv("JSEARCH_RAPIDAPI_HOST", "jsearch.p.rapidapi.com")
JSEARCH_SEARCH_URL = f"https://{JSEARCH_RAPIDAPI_HOST}/search-v2"
SEARCH_LOCATION = os.getenv("SEARCH_LOCATION", "Springfield, ST")
SEARCH_RADIUS_MILES = int(os.getenv("SEARCH_RADIUS_MILES", "50"))
JSEARCH_ALLOWED_PUBLISHERS = [
    p.strip() for p in os.getenv(
        "JSEARCH_ALLOWED_PUBLISHERS",
        "Indeed,Glassdoor,ZipRecruiter,CareerBuilder,Monster",
    ).split(",") if p.strip()
]
JSEARCH_FILTER_LOG_PATH = os.getenv("JSEARCH_FILTER_LOG_PATH", "/data/jsearch-publisher-filter.log")

TARGET_TITLES = [
    "Network Administrator",
    "Systems Administrator",
    "IT Administrator",
    "Endpoint Engineer",
    "Desktop Engineer",
    "M365 Engineer",
]
STATUSES = ["New", "Saved", "Applied", "Interview", "Rejected", "Offer"]


class StatusIn(BaseModel):
    status: str


class NoteIn(BaseModel):
    notes: str = ""


class GenerateIn(BaseModel):
    kind: str


class SettingsIn(BaseModel):
    notification_window: str = "24h"
    telegram_enabled: bool = True
    email_enabled: bool = False
    smtp_server: str = ""
    smtp_port: int = 587
    smtp_username: str = ""
    smtp_password: str = ""
    smtp_to: str = ""
    max_alerts_per_cycle: str = "10"
    quiet_hours_start: str = ""
    quiet_hours_end: str = ""


class ManualJob(BaseModel):
    source: str = "Manual"
    title: str
    company: str = ""
    location: str = ""
    salary: str = ""
    url: str = ""
    description: str = ""
    remote_type: str = "unknown"


def utcnow() -> str:
    return datetime.now(timezone.utc).isoformat()


def connect():
    Path(DB_PATH).parent.mkdir(parents=True, exist_ok=True)
    c = sqlite3.connect(DB_PATH)
    c.row_factory = sqlite3.Row
    return c


def init_db():
    with connect() as c:
        c.executescript(
            """
            CREATE TABLE IF NOT EXISTS jobs(
              id INTEGER PRIMARY KEY AUTOINCREMENT,
              source TEXT NOT NULL,
              publisher TEXT,
              title TEXT NOT NULL,
              company TEXT,
              location TEXT,
              salary TEXT,
              url TEXT,
              description TEXT,
              remote_type TEXT DEFAULT 'unknown',
              status TEXT DEFAULT 'New',
              match_score INTEGER DEFAULT 0,
              fit_summary TEXT DEFAULT '',
              score_breakdown TEXT DEFAULT '{}',
              ideal_match INTEGER DEFAULT 0,
              fingerprint TEXT UNIQUE,
              first_seen TEXT,
              last_seen TEXT,
              posted_at TEXT,
              alerted INTEGER DEFAULT 0,
              notes TEXT DEFAULT ''
            );
            CREATE TABLE IF NOT EXISTS generated_docs(
              id INTEGER PRIMARY KEY AUTOINCREMENT,
              job_id INTEGER,
              kind TEXT,
              title TEXT,
              content TEXT,
              created_at TEXT
            );
            CREATE TABLE IF NOT EXISTS settings(
              key TEXT PRIMARY KEY,
              value TEXT NOT NULL
            );
            CREATE INDEX IF NOT EXISTS idx_jobs_status ON jobs(status);
            CREATE INDEX IF NOT EXISTS idx_jobs_seen ON jobs(first_seen);
            CREATE INDEX IF NOT EXISTS idx_jobs_score ON jobs(match_score);
            CREATE INDEX IF NOT EXISTS idx_jobs_url ON jobs(url);
            CREATE INDEX IF NOT EXISTS idx_jobs_exact_duplicate ON jobs(title, company, location);
            """
        )
        cols = {r['name'] for r in c.execute('PRAGMA table_info(jobs)').fetchall()}
        if 'posted_at' not in cols:
            c.execute('ALTER TABLE jobs ADD COLUMN posted_at TEXT')
        if 'publisher' not in cols:
            c.execute('ALTER TABLE jobs ADD COLUMN publisher TEXT')
        c.execute('CREATE INDEX IF NOT EXISTS idx_jobs_posted ON jobs(posted_at)')
        c.execute('CREATE INDEX IF NOT EXISTS idx_jobs_publisher ON jobs(publisher)')
        # Existing rows predate the strict first-insert notification contract.
        # Mark them notified at startup so deploy/restart cannot resurrect old jobs as alerts.
        c.execute("UPDATE jobs SET alerted=1 WHERE alerted=0 AND first_seen IS NOT NULL")


def dedupe_jobs() -> int:
    """Remove duplicate jobs, keeping the newest row. Duplicate means same URL or same title/company/location."""
    with connect() as c:
        rows = c.execute(
            "SELECT * FROM jobs ORDER BY COALESCE(last_seen, first_seen, '') DESC, COALESCE(first_seen, '') DESC, id DESC"
        ).fetchall()
        seen_urls: set[str] = set()
        seen_identities: set[tuple[str, str, str]] = set()
        delete_ids = []
        for row in rows:
            url_key = canonical_url(row["url"])
            id_key = identity_key(row)
            duplicate = bool(url_key and url_key in seen_urls) or id_key in seen_identities
            if duplicate:
                delete_ids.append(row["id"])
                continue
            if url_key:
                seen_urls.add(url_key)
            seen_identities.add(id_key)
        if delete_ids:
            placeholders = ",".join("?" for _ in delete_ids)
            c.execute(f"DELETE FROM jobs WHERE id IN ({placeholders})", delete_ids)
        removed = len(delete_ids)
        if removed:
            print(f"dedupe removed {removed} duplicate job rows", flush=True)
        return removed


def backfill_linkedin_companies() -> int:
    """Best-effort repair for old LinkedIn rows scraped before company extraction existed."""
    with connect() as c:
        rows = c.execute(
            "SELECT id, url FROM jobs WHERE source='LinkedIn' AND COALESCE(company, '')='' AND COALESCE(url, '')!=''"
        ).fetchall()
        fixed = 0
        for row in rows:
            company = linkedin_company_from_url(row["url"])
            if company:
                c.execute("UPDATE jobs SET company=? WHERE id=?", (company, row["id"]))
                fixed += 1
        malformed = c.execute(
            "SELECT id, title, location FROM jobs WHERE source='LinkedIn' AND COALESCE(company, '')='' AND title LIKE '%' || char(10) || '%'"
        ).fetchall()
        for row in malformed:
            lines = [x.strip() for x in (row["title"] or "").splitlines() if x.strip()]
            if len(lines) >= 2:
                title, company = lines[0], clean_company(lines[1])
                location = row["location"] or (lines[2] if len(lines) >= 3 else "")
                if company:
                    c.execute("UPDATE jobs SET title=?, company=?, location=COALESCE(NULLIF(?, ''), location) WHERE id=?", (title, company, location, row["id"]))
                    fixed += 1
        if fixed:
            print(f"backfilled {fixed} LinkedIn company names", flush=True)
        return fixed


def parse_dt(value: str):
    if not value:
        return None
    try:
        return datetime.fromisoformat(str(value).replace("Z", "+00:00"))
    except Exception:
        return None


def posted_label(value: str, prefix: str = "Posted") -> str:
    dt = parse_dt(value)
    if not dt:
        return "Date unknown"
    if dt.tzinfo is None:
        dt = dt.replace(tzinfo=timezone.utc)
    now_dt = datetime.now(timezone.utc)
    delta = now_dt - dt.astimezone(timezone.utc)
    if delta.total_seconds() < 3600:
        mins = max(1, int(delta.total_seconds() // 60))
        return f"{prefix} {mins} minute" + ("" if mins == 1 else "s") + " ago"
    if delta.days < 1:
        hours = max(1, int(delta.total_seconds() // 3600))
        return f"{prefix} {hours} hour" + ("" if hours == 1 else "s") + " ago"
    if delta.days <= 14:
        return f"{prefix} {delta.days} day" + ("" if delta.days == 1 else "s") + " ago"
    fmt = "%b %-d" if dt.year == now_dt.year else "%b %-d, %Y"
    return f"{prefix} " + dt.astimezone(timezone.utc).strftime(fmt)


def rowdict(row):
    d = dict(row)
    try:
        d["score_breakdown"] = json.loads(d.get("score_breakdown") or "{}")
    except Exception:
        d["score_breakdown"] = {}
    d["ideal_match"] = bool(d.get("ideal_match"))
    actual_posted = d.get("posted_at") or ""
    found = d.get("first_seen") or d.get("last_seen") or ""
    d["posted_at"] = actual_posted
    d["posted_label"] = posted_label(actual_posted, "Posted") if actual_posted else posted_label(found, "Posted")
    return d


DEFAULT_SETTINGS = {
    "notification_window": f"{NOTIFICATION_WINDOW_HOURS}h",
    "telegram_enabled": "true",
    "email_enabled": "true" if SMTP_ENABLED else "false",
    "smtp_server": SMTP_SERVER,
    "smtp_port": str(SMTP_PORT),
    "smtp_username": SMTP_USERNAME,
    "smtp_password": SMTP_PASSWORD,
    "smtp_to": SMTP_TO,
    "max_alerts_per_cycle": str(ALERT_BATCH_LIMIT),
    "quiet_hours_start": QUIET_HOURS_START,
    "quiet_hours_end": QUIET_HOURS_END,
}


def settings_dict(redact: bool = True) -> dict:
    try:
        with connect() as c:
            rows = {r["key"]: r["value"] for r in c.execute("SELECT key,value FROM settings")}
    except Exception:
        rows = {}
    out = {**DEFAULT_SETTINGS, **rows}
    out["telegram_configured"] = bool(TELEGRAM_BOT_TOKEN and TELEGRAM_CHAT_ID)
    if redact and out.get("smtp_password"):
        out["smtp_password"] = "configured"
    return out


def save_settings(payload: dict):
    clean = {}
    for k, v in payload.items():
        if k not in DEFAULT_SETTINGS:
            continue
        if k == "smtp_password" and (not v or v == "configured"):
            continue
        clean[k] = str(v).strip() if not isinstance(v, bool) else ("true" if v else "false")
    with connect() as c:
        for k, v in clean.items():
            c.execute("INSERT INTO settings(key,value) VALUES(?,?) ON CONFLICT(key) DO UPDATE SET value=excluded.value", (k, v))
    return settings_dict()


def setting_bool(settings: dict, key: str) -> bool:
    return str(settings.get(key, "")).lower() in {"1", "true", "yes", "on"}


def setting_window_hours(value: str, default=24) -> int:
    v = str(value or "").lower().strip()
    if v in {"24h", "last 24 hours", "1"}: return 24
    if v in {"3d", "last 3 days", "3"}: return 72
    if v in {"7d", "last 7 days", "7"}: return 168
    try: return max(1, int(v.rstrip("h")))
    except Exception: return default


def setting_get(key: str, default: str = "") -> str:
    try:
        with connect() as c:
            row = c.execute("SELECT value FROM settings WHERE key=?", (key,)).fetchone()
            return row["value"] if row else default
    except Exception:
        return default


def setting_put(key: str, value: str) -> None:
    with connect() as c:
        c.execute("INSERT INTO settings(key,value) VALUES(?,?) ON CONFLICT(key) DO UPDATE SET value=excluded.value", (key, str(value)))


def jsearch_month_key() -> str:
    return datetime.now(timezone.utc).strftime("%Y-%m")


def jsearch_usage() -> dict:
    month = jsearch_month_key()
    saved_month = setting_get("jsearch_usage_month", month)
    if saved_month != month:
        return {"month": month, "used": 0, "limit": JSEARCH_MONTHLY_LIMIT, "remaining": JSEARCH_MONTHLY_LIMIT}
    try:
        used = max(0, int(setting_get("jsearch_usage_count", "0")))
    except Exception:
        used = 0
    return {"month": month, "used": used, "limit": JSEARCH_MONTHLY_LIMIT, "remaining": max(0, JSEARCH_MONTHLY_LIMIT - used)}


def jsearch_can_request() -> bool:
    if not (JSEARCH_ENABLED and JSEARCH_RAPIDAPI_KEY):
        return False
    return jsearch_usage()["remaining"] > 0


def jsearch_record_request(count: int = 1) -> None:
    usage = jsearch_usage()
    setting_put("jsearch_usage_month", usage["month"])
    setting_put("jsearch_usage_count", str(usage["used"] + max(0, count)))


def in_quiet_hours(settings: dict) -> bool:
    start, end = settings.get("quiet_hours_start", ""), settings.get("quiet_hours_end", "")
    if not start or not end:
        return False
    now = datetime.now().strftime("%H:%M")
    return start <= now < end if start <= end else now >= start or now < end


def max_alerts(settings: dict) -> int | None:
    value = str(settings.get("max_alerts_per_cycle", "10")).lower()
    if value in {"unlimited", "0", "none"}: return None
    try: return max(1, int(value))
    except Exception: return ALERT_BATCH_LIMIT


def norm(text) -> str:
    return re.sub(r"\s+", " ", str(text or "").strip().lower())


def canonical_url(url: str) -> str:
    raw = (url or "").strip()
    if not raw:
        return ""
    try:
        parsed = urllib.parse.urlsplit(raw)
        if not parsed.scheme or not parsed.netloc:
            return raw.split("?")[0].split("#")[0].rstrip("/").lower()
        path = re.sub(r"/+", "/", urllib.parse.unquote(parsed.path or "")).rstrip("/")
        return urllib.parse.urlunsplit((parsed.scheme.lower(), parsed.netloc.lower(), path, "", ""))
    except Exception:
        return raw.split("?")[0].split("#")[0].rstrip("/").lower()


def field(job, name: str, default=""):
    if hasattr(job, "get"):
        return job.get(name, default)
    try:
        return job[name]
    except Exception:
        return default


def identity_key(job: dict) -> tuple[str, str, str]:
    return (norm(field(job, "title")), norm(field(job, "company")), norm(field(job, "location")))


def fingerprint(job: dict) -> str:
    url = canonical_url(job.get("url", ""))
    if url:
        return hashlib.sha256(("url:" + url).encode()).hexdigest()
    title, company, location = identity_key(job)
    return hashlib.sha256(("tcl:" + title + "|" + company + "|" + location).encode()).hexdigest()


def find_existing_job(c: sqlite3.Connection, job: dict, fp: str):
    old = c.execute("SELECT * FROM jobs WHERE fingerprint=?", (fp,)).fetchone()
    if old:
        return old

    url = canonical_url(job.get("url", ""))
    if url:
        for row in c.execute("SELECT * FROM jobs WHERE COALESCE(url, '') != ''"):
            if canonical_url(row["url"]) == url:
                return row

    key = identity_key(job)
    for row in c.execute("SELECT * FROM jobs"):
        if identity_key(row) == key:
            return row
    return None


def salary_floor(text: str) -> int:
    nums = []
    for match in re.finditer(r"\$?\s*(\d{2,3})(?:[,\.]?(\d{3}))?\s*([kK])?", text or ""):
        n = int(match.group(1) + (match.group(2) or ""))
        if match.group(3) or n < 1000:
            n *= 1000
        nums.append(n)
    return max(nums) if nums else 0


def infer_remote_type(title="", location="", description="") -> str:
    blob = " ".join([title or "", location or "", description or ""]).lower()
    if "hybrid" in blob:
        return "hybrid"
    if "remote" in blob or "work from home" in blob:
        return "remote"
    if "onsite" in blob or "on-site" in blob or "on site" in blob:
        return "onsite"
    return "unknown"


def estimate_distance(location: str) -> int:
    loc = (location or "").lower()
    if "springfield" in loc:
        return 0
    near = {
        "chelmsford": 5,
        "billerica": 10,
        "andover": 15,
        "bedford": 17,
        "nashua": 17,
        "woburn": 22,
        "cambridge": 28,
        "waltham": 28,
        "boston": 31,
        "manchester": 38,
        "worcester": 45,
    }
    for city, miles in near.items():
        if city in loc:
            return miles
    if "remote" in loc:
        return 999
    if " ma" in " " + loc or "massachusetts" in loc or " nh" in " " + loc or "new hampshire" in loc:
        return 50
    return 999


def heuristic_score(job: dict):
    title = job.get("title", "")
    desc = job.get("description", "")
    loc = job.get("location", "")
    salary = job.get("salary", "")
    blob = " ".join([title, desc, loc, salary]).lower()
    title_l = title.lower()
    score = 18
    bits = {}

    title_weights = {
        "senior network administrator": 26,
        "network administrator": 24,
        "systems administrator": 21,
        "system administrator": 21,
        "it administrator": 20,
        "endpoint engineer": 22,
        "desktop engineer": 17,
        "m365 engineer": 18,
        "azure administrator": 18,
        "help desk": -10,
        "technician": -6,
    }
    bits["title_match"] = max([weight for phrase, weight in title_weights.items() if phrase in title_l] or [6 if any(t.lower() in title_l for t in TARGET_TITLES) else 0])
    score += bits["title_match"]

    skill_terms = {
        "azure": 4, "intune": 5, "endpoint": 4, "m365": 4, "microsoft 365": 4,
        "active directory": 4, "entra": 4, "network": 3, "firewall": 3, "switch": 2,
        "routing": 2, "hyper-v": 4, "vmware": 2, "security": 2, "powershell": 3,
        "sccm": 3, "autopilot": 4, "windows server": 3,
    }
    hits = {term: weight for term, weight in skill_terms.items() if term in blob}
    bits["skill_alignment"] = min(28, sum(hits.values()))
    bits["skill_hits"] = sorted(hits)
    score += bits["skill_alignment"]

    seniority = 8 if any(x in blob for x in ["senior", "lead", "sr.", "sr "]) else 4 if "administrator" in title_l or "engineer" in title_l else 0
    bits["seniority"] = seniority
    score += seniority

    work_mode = job.get("remote_type") or infer_remote_type(title, loc, desc)
    distance = estimate_distance(loc)
    loc_score = 0
    if work_mode == "hybrid" and distance <= 40:
        loc_score = 22
    elif work_mode == "remote":
        loc_score = 17
    elif distance <= 20:
        loc_score = 14
    elif distance <= 50:
        loc_score = 8
    elif work_mode == "onsite":
        loc_score = -8
    bits["location_work_mode"] = loc_score
    bits["estimated_distance_miles"] = distance
    bits["work_mode"] = work_mode
    score += loc_score

    sf = salary_floor(salary)
    if sf >= 120000:
        sal_score = 16
    elif sf >= 115000:
        sal_score = 14
    elif sf >= 105000:
        sal_score = 7
    elif sf >= 90000:
        sal_score = 1
    elif sf:
        sal_score = -12
    else:
        sal_score = -2
    bits["salary"] = sal_score
    bits["salary_floor"] = sf
    score += sal_score

    penalties = 0
    if any(x in blob for x in ["entry level", "junior", "intern", "contract only", "temporary"]):
        penalties -= 10
    if "clearance" in blob and "secret" in blob:
        penalties -= 4
    bits["penalties"] = penalties
    score += penalties

    # Small deterministic differentiator so sparse scraped cards do not all become the same 73/100 clone.
    fp_seed = json.dumps({"title": title, "company": job.get("company", ""), "location": loc, "salary": salary, "source": job.get("source", "")}, sort_keys=True)
    bits["specificity_adjustment"] = (int(hashlib.sha256(fp_seed.encode()).hexdigest()[:2], 16) % 7) - 3
    score += bits["specificity_adjustment"]

    score = max(0, min(100, int(round(score))))
    ideal = work_mode == "hybrid" and distance <= 40 and sf >= 115000 and bits["title_match"] >= 20
    summary = "Strong fit" if score >= 80 else "Possible fit" if score >= 60 else "Below target"
    if sf and sf < 115000:
        summary += "; salary appears below target"
    if not sf:
        summary += "; salary not listed"
    if work_mode not in ["remote", "hybrid"] and distance > 40:
        summary += "; location/work mode may be weak"
    return score, summary, bits, ideal


def claude_json(prompt: str) -> Optional[dict]:
    if not ANTHROPIC_API_KEY:
        return None
    try:
        r = requests.post(
            "https://api.anthropic.com/v1/messages",
            headers={"x-api-key": ANTHROPIC_API_KEY, "anthropic-version": "2023-06-01", "content-type": "application/json"},
            json={"model": ANTHROPIC_MODEL, "max_tokens": 1400, "messages": [{"role": "user", "content": prompt}]},
            timeout=45,
        )
        r.raise_for_status()
        text = "\n".join(part.get("text", "") for part in r.json().get("content", []))
        m = re.search(r"\{.*\}", text, re.S)
        return json.loads(m.group(0)) if m else None
    except Exception:
        return None


def score_job(job: dict):
    score, summary, breakdown, ideal = heuristic_score(job)
    prompt = (
        "Return JSON only with keys score integer 0-100, summary string, breakdown object, ideal_match boolean. "
        "Score this job for Michael Dziegiel: Senior Network Administrator, 20+ years, Azure, Intune, Hyper-V, endpoint management, SCCM, M365, networking, security, Hyper-V. "
        "Target salary is configurable. Ideal is hybrid within the configured commute radius, at/above target salary, senior infrastructure/endpoint/network role. "
        "Use the specific title, location, salary, work mode, and skill overlap. Do not return a generic middle score. "
        "Jobs must produce different scores when title/location/salary/skills differ. Job: " + json.dumps(job, sort_keys=True)[:7000]
    )
    ai = claude_json(prompt)
    if ai:
        try:
            ai_score = int(max(0, min(100, ai.get("score", score))))
            # Claude has been flattening sparse jobs to the same value. Blend with the deterministic profile score
            # so actual title/location/salary differences survive instead of becoming 73/100 wallpaper.
            score = int(round((score * 0.65) + (ai_score * 0.35)))
        except Exception:
            pass
        summary = str(ai.get("summary") or summary)[:500]
        if isinstance(ai.get("breakdown"), dict):
            breakdown = {**breakdown, "claude": ai["breakdown"]}
        ideal = bool(ai.get("ideal_match", ideal)) and ideal
    return score, summary, breakdown, ideal


def upsert_job(job: dict):
    job = {k: (v or "") for k, v in job.items()}
    if not job.get("remote_type") or job.get("remote_type") == "unknown":
        job["remote_type"] = infer_remote_type(job.get("title"), job.get("location"), job.get("description"))
    fp = fingerprint(job)
    score, summary, breakdown, ideal = score_job(job)
    with connect() as c:
        old = find_existing_job(c, job, fp)
        if old:
            c.execute(
                "UPDATE jobs SET last_seen=?, posted_at=COALESCE(NULLIF(?, ''), posted_at), publisher=COALESCE(NULLIF(?, ''), publisher), company=COALESCE(NULLIF(company, ''), NULLIF(?, ''), company), location=COALESCE(NULLIF(location, ''), NULLIF(?, ''), location), salary=COALESCE(NULLIF(?, ''), salary), url=COALESCE(NULLIF(?, ''), url), description=COALESCE(NULLIF(?, ''), description), match_score=?, fit_summary=?, score_breakdown=?, ideal_match=?, fingerprint=COALESCE(NULLIF(fingerprint, ''), ?) WHERE id=?",
                (utcnow(), job.get("posted_at", ""), job.get("publisher", ""), job.get("company", ""), job.get("location", ""), job.get("salary", ""), job.get("url", ""), job.get("description", ""), score, summary, json.dumps(breakdown), int(ideal), fp, old["id"]),
            )
            return rowdict(c.execute("SELECT * FROM jobs WHERE id=?", (old["id"],)).fetchone()), False
        c.execute(
            "INSERT INTO jobs(source,publisher,title,company,location,salary,url,description,remote_type,match_score,fit_summary,score_breakdown,ideal_match,fingerprint,first_seen,last_seen,posted_at,alerted) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,0)",
            (job.get("source"), job.get("publisher", ""), job.get("title"), job.get("company"), job.get("location"), job.get("salary"), job.get("url"), job.get("description"), job.get("remote_type"), score, summary, json.dumps(breakdown), int(ideal), fp, utcnow(), utcnow(), job.get("posted_at", "")),
        )
        rid = c.execute("SELECT last_insert_rowid()").fetchone()[0]
        return rowdict(c.execute("SELECT * FROM jobs WHERE id=?", (rid,)).fetchone()), True


def clean_text(text: str) -> str:
    return re.sub(r"\s+", " ", html.unescape(re.sub(r"<[^>]+>", " ", text or ""))).strip()


def clean_company(text: str) -> str:
    company = clean_text(text)
    company = re.sub(r"\s*(?:logo|company logo)$", "", company, flags=re.I).strip(" -·|\t\n\r")
    noise = {"", "promoted", "be an early applicant", "actively hiring", "view job"}
    return "" if company.lower() in noise else company[:160]


def linkedin_company_from_url(url: str) -> str:
    """Extract company from LinkedIn public slugs like title-at-acme-corp-123456."""
    path = urllib.parse.unquote(urllib.parse.urlsplit(url or "").path).rstrip("/")
    slug = path.rsplit("/", 1)[-1]
    m = re.search(r"-at-(?P<company>.+?)(?:-\d+)?$", slug, re.I)
    if not m:
        return ""
    company = re.sub(r"[^\w\s&.+-]+", " ", m.group("company").replace("-", " "), flags=re.UNICODE)
    return clean_company(company.title())


def html_attr(tag: str, attr: str):
    m = re.search(rf'{attr}=["\']([^"\']+)["\']', tag or "", re.I)
    return html.unescape(m.group(1)) if m else ""


def extract_linkedin_company(card_html: str) -> str:
    patterns = [
        r'<h4[^>]*class=["\'][^"\']*base-search-card__subtitle[^"\']*["\'][^>]*>(?P<value>.*?)</h4>',
        r'<a[^>]*class=["\'][^"\']*hidden-nested-link[^"\']*["\'][^>]*>(?P<value>.*?)</a>',
        r'<a[^>]*data-tracking-control-name=["\'][^"\']*job_card_company[^"\']*["\'][^>]*>(?P<value>.*?)</a>',
        r'<span[^>]*class=["\'][^"\']*job-card-container__primary-description[^"\']*["\'][^>]*>(?P<value>.*?)</span>',
    ]
    for pattern in patterns:
        m = re.search(pattern, card_html or "", re.S | re.I)
        if m:
            company = clean_company(m.group("value"))
            if company:
                return company
    alt = re.search(r'<img[^>]+alt=["\'](?P<value>[^"\']+)["\'][^>]*>', card_html or "", re.I)
    if alt:
        company = clean_company(alt.group("value"))
        if company:
            return company
    return ""


def extract_linkedin_title(card_html: str, fallback: str) -> str:
    for pattern in [
        r'<h3[^>]*class=["\'][^"\']*base-search-card__title[^"\']*["\'][^>]*>(?P<value>.*?)</h3>',
        r'<a[^>]+href=["\']https://www\.linkedin\.com/jobs/view/[^"\']+["\'][^>]*>(?P<value>.*?)</a>',
    ]:
        m = re.search(pattern, card_html or "", re.S | re.I)
        if m:
            title = clean_text(m.group("value"))
            if title:
                return title
    return fallback


def extract_linkedin_location(card_html: str) -> str:
    m = re.search(r'<span[^>]*class=["\'][^"\']*job-search-card__location[^"\']*["\'][^>]*>(?P<value>.*?)</span>', card_html or "", re.S | re.I)
    return clean_text(m.group("value")) if m else f"{SEARCH_LOCATION} / Remote"



def relative_posted_at(text: str) -> str:
    raw = clean_text(text).lower()
    m = re.search(r'(\d+)\s*(minute|minutes|hour|hours|day|days|week|weeks|month|months)\s+ago', raw)
    if not m:
        if 'just now' in raw or 'today' in raw:
            return utcnow()
        return ''
    n = int(m.group(1)); unit = m.group(2)
    days = 0
    if unit.startswith('minute'):
        return (datetime.now(timezone.utc) - timedelta(minutes=n)).isoformat()
    if unit.startswith('hour'):
        return (datetime.now(timezone.utc) - timedelta(hours=n)).isoformat()
    if unit.startswith('day'):
        days = n
    elif unit.startswith('week'):
        days = n * 7
    elif unit.startswith('month'):
        days = n * 30
    return (datetime.now(timezone.utc) - timedelta(days=days)).isoformat()


def extract_linkedin_posted_at(card_html: str) -> str:
    for tag in re.findall(r'<time[^>]*>.*?</time>', card_html or '', re.S | re.I):
        dt = html_attr(tag, 'datetime')
        if dt:
            parsed = parse_dt(dt) or parse_dt(dt + 'T00:00:00+00:00')
            if parsed:
                if parsed.tzinfo is None:
                    parsed = parsed.replace(tzinfo=timezone.utc)
                return parsed.astimezone(timezone.utc).isoformat()
        rel = relative_posted_at(tag)
        if rel:
            return rel
    for pattern in [r'(\d+\s+(?:minute|minutes|hour|hours|day|days|week|weeks|month|months)\s+ago)', r'(just now|today)']:
        m = re.search(pattern, card_html or '', re.I)
        if m:
            rel = relative_posted_at(m.group(1))
            if rel:
                return rel
    return ''

def parse_linkedin_jobs(page_html: str, fallback_title: str) -> list[dict]:
    jobs = []
    seen = set()
    cards = re.findall(r'<(?:li|div)[^>]*(?:base-card|job-search-card)[^>]*>.*?</(?:li|div)>', page_html or "", re.S | re.I)
    if not cards:
        cards = [m.group(0) for m in re.finditer(r'.{0,2500}https://www\.linkedin\.com/jobs/view/[^"#?<]+.{0,2500}', page_html or "", re.S | re.I)]
    for card in cards:
        url_m = re.search(r'https://www\.linkedin\.com/jobs/view/[^?"#<\s]+', card, re.I)
        if not url_m:
            continue
        url = html.unescape(url_m.group(0)).rstrip('/')
        if url in seen:
            continue
        seen.add(url)
        jobs.append({
            "source": "LinkedIn",
            "title": extract_linkedin_title(card, fallback_title),
            "company": extract_linkedin_company(card) or linkedin_company_from_url(url),
            "location": extract_linkedin_location(card),
            "salary": "",
            "url": url,
            "description": "LinkedIn public listing",
            "remote_type": "unknown",
            "posted_at": extract_linkedin_posted_at(card),
        })
    return jobs


def fetch(url: str) -> str:
    return requests.get(url, headers={"User-Agent": "Mozilla/5.0 JobWatchAssistant/1.0", "Accept-Language": "en-US,en;q=0.9"}, timeout=20).text


def jsearch_salary(job: dict) -> str:
    min_salary = job.get("job_min_salary")
    max_salary = job.get("job_max_salary")
    currency = job.get("job_salary_currency") or "USD"
    period = job.get("job_salary_period") or "year"
    if not min_salary and not max_salary:
        return job.get("job_salary") or ""
    def fmt(value):
        try:
            return f"${int(float(value)):,}"
        except Exception:
            return str(value)
    if min_salary and max_salary and min_salary != max_salary:
        salary = f"{fmt(min_salary)} - {fmt(max_salary)}"
    else:
        salary = fmt(min_salary or max_salary)
    if currency and currency != "USD":
        salary += f" {currency}"
    if period:
        salary += f" / {period}"
    return salary


def jsearch_location(job: dict) -> str:
    parts = [job.get("job_city"), job.get("job_state"), job.get("job_country")]
    location = ", ".join(str(part).strip() for part in parts if part)
    return location or job.get("job_location") or ("Remote" if job.get("job_is_remote") else "")


def parse_jsearch_datetime(value: str) -> str:
    if not value:
        return ""
    try:
        dt = datetime.fromisoformat(str(value).replace("Z", "+00:00"))
        if dt.tzinfo is None:
            dt = dt.replace(tzinfo=timezone.utc)
        return dt.astimezone(timezone.utc).isoformat()
    except Exception:
        return ""


def normalize_jsearch_job(job: dict, fallback_title: str = "") -> dict:
    title = clean_text(job.get("job_title") or fallback_title)
    company = clean_company(job.get("employer_name") or "")
    location = clean_text(jsearch_location(job))
    description = clean_text(job.get("job_description") or "JSearch API listing")
    url = job.get("job_apply_link") or job.get("job_google_link") or ""
    remote_type = "remote" if job.get("job_is_remote") else infer_remote_type(title, location, description)
    return {
        "source": "jsearch",
        "publisher": clean_text(job.get("job_publisher") or ""),
        "title": title,
        "company": company,
        "location": location,
        "salary": jsearch_salary(job),
        "url": url,
        "description": description,
        "remote_type": remote_type,
        "posted_at": parse_jsearch_datetime(job.get("job_posted_at_datetime_utc") or ""),
    }


def jsearch_publisher_allowed(publisher: str) -> bool:
    allowed = {p.casefold() for p in JSEARCH_ALLOWED_PUBLISHERS}
    return clean_text(publisher).casefold() in allowed


def jsearch_filter_summary(stats: dict[str, dict[str, int]]) -> str:
    parts = []
    for publisher in sorted(stats, key=lambda p: p.casefold()):
        kept = stats[publisher].get("kept", 0)
        filtered = stats[publisher].get("filtered", 0)
        if kept:
            parts.append(f"{kept} {publisher} kept")
        if filtered:
            parts.append(f"{filtered} {publisher} filtered")
    return ", ".join(parts) if parts else "no publisher stats"


def log_jsearch_filter_breakdown(title: str, total: int, stats: dict[str, dict[str, int]]) -> None:
    message = f"{total} found: {jsearch_filter_summary(stats)}"
    line = f"{utcnow()} jsearch publisher filter title={title!r}: {message}"
    print(line, flush=True)
    try:
        path = Path(JSEARCH_FILTER_LOG_PATH)
        path.parent.mkdir(parents=True, exist_ok=True)
        with path.open("a", encoding="utf-8") as fh:
            fh.write(line + "\n")
    except Exception as exc:
        print(f"failed to write jsearch publisher filter log: {exc}", flush=True)


def scrape_jsearch(title: str):
    if not jsearch_can_request():
        return []
    headers = {"X-RapidAPI-Key": JSEARCH_RAPIDAPI_KEY, "X-RapidAPI-Host": JSEARCH_RAPIDAPI_HOST}
    params = {"query": f"{title} in {SEARCH_LOCATION}", "num_pages": "1", "country": "us", "date_posted": "week"}
    try:
        r = requests.get(JSEARCH_SEARCH_URL, headers=headers, params=params, timeout=25)
        jsearch_record_request(1)
        r.raise_for_status()
        payload = r.json()
        data = payload.get("data") if isinstance(payload, dict) else []
        rows = data.get("jobs", []) if isinstance(data, dict) else data
        jobs = []
        seen = set()
        stats: dict[str, dict[str, int]] = {}
        for item in rows or []:
            if not isinstance(item, dict):
                continue
            job = normalize_jsearch_job(item, title)
            publisher = job.get("publisher") or "Unknown"
            publisher_stats = stats.setdefault(publisher, {"kept": 0, "filtered": 0})
            if not jsearch_publisher_allowed(publisher):
                publisher_stats["filtered"] += 1
                continue
            key = identity_key(job)
            if not job.get("title") or key in seen:
                continue
            seen.add(key)
            publisher_stats["kept"] += 1
            jobs.append(job)
        log_jsearch_filter_breakdown(title, len(rows or []), stats)
        return jobs[:15]
    except Exception as exc:
        print(f"jsearch scrape failed for {title}: {exc}", flush=True)
        return []


def scrape_linkedin(title: str):
    q = urllib.parse.quote(title)
    loc = urllib.parse.quote(SEARCH_LOCATION)
    url = f"https://www.linkedin.com/jobs/search?keywords={q}&location={loc}&distance={SEARCH_RADIUS_MILES}"
    try:
        return parse_linkedin_jobs(fetch(url), title)[:15]
    except Exception:
        return []


def scrape_dice(title: str):
    q = urllib.parse.quote(title)
    loc = urllib.parse.quote(SEARCH_LOCATION)
    url = f"https://www.dice.com/jobs?q={q}&location={loc}&radius={SEARCH_RADIUS_MILES}&radiusUnit=mi&page=1&pageSize=20"
    jobs = []

    def dice_value(blob: str, key: str) -> str:
        m = re.search(r'\\"' + re.escape(key) + r'\\":\\"(?P<v>(?:\\\\.|[^"\\\\])*)\\"', blob, re.S)
        if not m:
            return ""
        try:
            return json.loads('"' + m.group("v").replace('"', r'\"') + '"')
        except Exception:
            return m.group("v").replace(r"\/", "/").replace(r"\u0026", "&")

    try:
        text = fetch(url)
        matches = list(re.finditer(r'\\"detailsPageUrl\\":\\"(?P<url>(?:\\\\.|[^"\\\\])*)\\"', text, re.S))
        for i, m in enumerate(matches):
            end = matches[i + 1].start() if i + 1 < len(matches) else m.end() + 5000
            blob = text[m.start():end]
            loc = ""
            loc_m = re.search(r'\\"jobLocation\\":\{(?P<loc>.*?)\}', blob, re.S)
            if loc_m:
                loc = dice_value(loc_m.group("loc"), "displayName") or dice_value(loc_m.group("loc"), "city")
            job = {
                "source": "Dice",
                "title": dice_value(blob, "title") or title,
                "company": dice_value(blob, "companyName"),
                "location": loc,
                "salary": dice_value(blob, "salary"),
                "url": dice_value(m.group(0), "detailsPageUrl"),
                "description": "Dice public listing",
                "remote_type": "unknown",
                "posted_at": relative_posted_at(dice_value(blob, "postedDate")) or dice_value(blob, "postedDate"),
            }
            if job["url"] and job["url"] not in {j["url"] for j in jobs}:
                jobs.append(job)
    except Exception:
        pass
    return jobs[:15]


def scrape_ziprecruiter(title: str):
    q = urllib.parse.quote(title)
    loc = urllib.parse.quote_plus(SEARCH_LOCATION)
    url = f"https://www.ziprecruiter.com/jobs-search?search={q}&location={loc}&radius={SEARCH_RADIUS_MILES}"
    jobs = []
    try:
        text = fetch(url)
        for m in re.finditer(r'<a[^>]+href="(?P<url>https://www\.ziprecruiter\.com/[^"#]+)"[^>]*>(?P<title>[^<]{5,120})</a>', text, re.S):
            jobs.append({"source": "ZipRecruiter", "title": clean_text(m.group("title")), "company": "", "location": f"{SEARCH_LOCATION} / Remote", "salary": "", "url": m.group("url"), "description": "ZipRecruiter public listing", "remote_type": "unknown"})
    except Exception:
        pass
    return jobs[:15]


def scrape_indeed_mcp(title: str):
    jobs = []
    try:
        payload = {"jsonrpc": "2.0", "id": 1, "method": "tools/call", "params": {"name": "search_jobs", "arguments": {"query": title, "location": SEARCH_LOCATION, "radius": SEARCH_RADIUS_MILES, "remote": True}}}
        r = requests.post(INDEED_MCP_URL, json=payload, headers={"Accept": "application/json"}, timeout=25)
        if r.ok:
            data = r.json()
            text = json.dumps(data.get("result", data))
            for obj in re.findall(r"\{[^{}]*(?:jobTitle|title)[^{}]*\}", text):
                try:
                    j = json.loads(obj)
                    jobs.append({"source": "Indeed", "title": j.get("jobTitle") or j.get("title") or title, "company": j.get("company") or j.get("companyName") or "", "location": j.get("location") or "", "salary": j.get("salary") or "", "url": j.get("url") or j.get("jobUrl") or "", "description": j.get("description") or "", "remote_type": "unknown"})
                except Exception:
                    pass
    except Exception:
        pass
    return jobs[:15]


def scrape_all():
    found = []
    for title in TARGET_TITLES:
        for fn in [scrape_indeed_mcp, scrape_linkedin, scrape_dice, scrape_ziprecruiter]:
            found.extend(fn(title))
        for fn in [scrape_linkedin, scrape_dice, scrape_ziprecruiter]:
            found.extend(fn(title + " remote"))
    # JSearch is API-backed and quota-limited. Rotate target titles instead of spending one
    # RapidAPI request per title every six hours. The free tier is 200 requests/month.
    if JSEARCH_ENABLED and JSEARCH_RAPIDAPI_KEY and JSEARCH_PER_SCRAPE_LIMIT > 0:
        usage = jsearch_usage()
        start = usage["used"] % len(TARGET_TITLES)
        titles = TARGET_TITLES[start:] + TARGET_TITLES[:start]
        for title in titles[:JSEARCH_PER_SCRAPE_LIMIT]:
            found.extend(scrape_jsearch(title))
    new_jobs = []
    for raw in found:
        if not raw.get("title"):
            continue
        saved, is_new = upsert_job(raw)
        if is_new:
            new_jobs.append(saved)
    alert_new_jobs(new_jobs)
    return {"found": len(found), "new": len(new_jobs)}


def job_alert_text(job: dict) -> str:
    return f"New job: {job['title']}\n{job.get('company','')} · {job.get('location','')} · {job.get('salary','')}\nScore: {job.get('match_score')} · {job.get('fit_summary')}\nDate: {job.get('posted_label') or (posted_label(job.get('posted_at'), 'Posted') if job.get('posted_at') else posted_label(job.get('first_seen'), 'Found'))}\n{job.get('url','')}"


def send_telegram(job: dict):
    if not (TELEGRAM_BOT_TOKEN and TELEGRAM_CHAT_ID):
        return
    try:
        requests.post(f"https://api.telegram.org/bot{TELEGRAM_BOT_TOKEN}/sendMessage", json={"chat_id": TELEGRAM_CHAT_ID, "text": job_alert_text(job), "disable_web_page_preview": False}, timeout=10)
    except Exception:
        pass


def send_telegram_summary(count: int):
    if not (TELEGRAM_BOT_TOKEN and TELEGRAM_CHAT_ID):
        return
    try:
        text = f"{count} new jobs found — check the dashboard"
        requests.post(f"https://api.telegram.org/bot{TELEGRAM_BOT_TOKEN}/sendMessage", json={"chat_id": TELEGRAM_CHAT_ID, "text": f"{text}\n{DASHBOARD_URL}", "disable_web_page_preview": True}, timeout=10)
    except Exception:
        pass


def send_email(job: dict, settings: dict):
    if not setting_bool(settings, "email_enabled"):
        return
    server, to_addr = settings.get("smtp_server", ""), settings.get("smtp_to", "")
    if not server or not to_addr:
        return
    msg = EmailMessage()
    msg["Subject"] = f"Job Watch: {job.get('match_score')}/100 {job.get('title')}"
    msg["From"] = settings.get("smtp_username") or "job-watch@mrdtech.local"
    msg["To"] = to_addr
    msg.set_content(job_alert_text(job))
    with smtplib.SMTP(server, int(settings.get("smtp_port") or 587), timeout=20) as smtp:
        smtp.ehlo()
        try:
            smtp.starttls(); smtp.ehlo()
        except Exception:
            pass
        if settings.get("smtp_username"):
            smtp.login(settings.get("smtp_username"), settings.get("smtp_password", ""))
        smtp.send_message(msg)


def send_email_summary(count: int, settings: dict):
    if not setting_bool(settings, "email_enabled") or not settings.get("smtp_server") or not settings.get("smtp_to"):
        return
    msg = EmailMessage(); msg["Subject"] = f"Job Watch: {count} new jobs"; msg["From"] = settings.get("smtp_username") or "job-watch@mrdtech.local"; msg["To"] = settings.get("smtp_to")
    msg.set_content(f"{count} new jobs found.\n{DASHBOARD_URL}")
    with smtplib.SMTP(settings.get("smtp_server"), int(settings.get("smtp_port") or 587), timeout=20) as smtp:
        smtp.ehlo()
        try:
            smtp.starttls(); smtp.ehlo()
        except Exception:
            pass
        if settings.get("smtp_username"):
            smtp.login(settings.get("smtp_username"), settings.get("smtp_password", ""))
        smtp.send_message(msg)


def alert_new_jobs(jobs):
    ids = [int(j["id"]) for j in jobs if j.get("id")]
    if not ids:
        return
    settings = settings_dict(redact=False)
    placeholders = ",".join("?" for _ in ids)
    cutoff = (datetime.now(timezone.utc) - timedelta(hours=setting_window_hours(settings.get("notification_window")))).isoformat()
    with connect() as c:
        unalerted = [rowdict(r) for r in c.execute(f"SELECT * FROM jobs WHERE id IN ({placeholders}) AND alerted=0 AND COALESCE(NULLIF(posted_at,''),first_seen,last_seen)>=? ORDER BY match_score DESC, first_seen DESC", [*ids, cutoff]).fetchall()]
        limit = max_alerts(settings)
        if in_quiet_hours(settings):
            c.execute(f"UPDATE jobs SET alerted=1 WHERE id IN ({placeholders})", ids)
            return
        if limit is not None and len(unalerted) > limit:
            if setting_bool(settings, "telegram_enabled"):
                send_telegram_summary(len(unalerted))
            send_email_summary(len(unalerted), settings)
        else:
            for j in unalerted:
                if setting_bool(settings, "telegram_enabled"):
                    send_telegram(j)
                try:
                    send_email(j, settings)
                except Exception:
                    pass
        c.execute(f"UPDATE jobs SET alerted=1 WHERE id IN ({placeholders})", ids)


def load_resume_text() -> str:
    p = Path(RESUME_BUILDER_DB)
    fallback = "Michael Dziegiel: Senior Network Administrator, 20+ years, Azure, Intune, Hyper-V, endpoint, Microsoft 365, Active Directory, networking, security."
    if not p.exists():
        return fallback
    try:
        rc = sqlite3.connect(str(p))
        rc.row_factory = sqlite3.Row
        row = rc.execute("SELECT data_json FROM resumes ORDER BY updated_at DESC LIMIT 1").fetchone()
        if not row:
            return fallback
        data = json.loads(row["data_json"])
        parts = [json.dumps(data.get("contact", {})), data.get("summary", "")]
        parts.extend(json.dumps(x) for x in data.get("experience", []))
        parts.append(json.dumps(data.get("skills", [])))
        return "\n".join(parts)[:10000]
    except Exception:
        return fallback


def generate_content(job: dict, kind: str) -> str:
    resume = load_resume_text()
    if kind == "cover_letter":
        fallback = f"Dear Hiring Manager,\n\nI am writing to express interest in the {job['title']} role at {job.get('company') or 'your organization'}. My background as a Senior Network Administrator with 20+ years across Microsoft 365, Azure, Intune, endpoint management, Hyper-V, networking, and security aligns well with the role requirements.\n\nI would welcome the opportunity to discuss how my infrastructure and endpoint experience can help your team deliver secure, reliable IT operations.\n\nSincerely,\nMichael Dziegiel"
        prompt = "Return JSON only as {\"content\":\"...\"}. Write a professional tailored cover letter for Michael Dziegiel. Resume/profile: " + resume + " Job: " + json.dumps(job)[:8000]
    else:
        fallback = f"Michael Dziegiel\nSenior Network Administrator\n\nSummary\nSenior Network Administrator with 20+ years of experience in Azure, Intune, Hyper-V, endpoint engineering, Microsoft 365, Active Directory, networking, and IT operations. Tailored target: {job['title']} at {job.get('company') or 'target employer'}.\n\nCore Skills\nAzure, Intune, Microsoft 365, Endpoint Management, Hyper-V, Active Directory, Networking, Security, Troubleshooting."
        prompt = "Return JSON only as {\"content\":\"...\"}. Create an ATS-optimized resume version for Michael Dziegiel. Keep it truthful. Resume/profile: " + resume + " Job: " + json.dumps(job)[:8000]
    out = claude_json(prompt)
    return (out or {}).get("content") or fallback


def rescore_existing_jobs(limit: int = 1000) -> int:
    updated = 0
    with connect() as c:
        rows = c.execute("SELECT * FROM jobs ORDER BY COALESCE(last_seen, first_seen, '') DESC LIMIT ?", (limit,)).fetchall()
        for row in rows:
            job = rowdict(row)
            # Existing-row backfills must be fast and deterministic. Do not call Claude for hundreds of rows at startup.
            score, summary, breakdown, ideal = heuristic_score(job)
            if score != row["match_score"] or row["match_score"] == 73:
                c.execute("UPDATE jobs SET match_score=?, fit_summary=?, score_breakdown=?, ideal_match=? WHERE id=?", (score, summary, json.dumps(breakdown), int(ideal), row["id"]))
                updated += 1
    if updated:
        print(f"rescored {updated} existing jobs", flush=True)
    return updated


async def cron_loop():
    await asyncio.sleep(8)
    while True:
        try:
            await asyncio.to_thread(scrape_all)
        except Exception as exc:
            print("scrape failed", exc, flush=True)
        await asyncio.sleep(6 * 60 * 60)


@asynccontextmanager
async def lifespan(app):
    init_db()
    dedupe_jobs()
    backfill_linkedin_companies()
    rescore_existing_jobs()
    task = asyncio.create_task(cron_loop()) if SCHEDULER_ENABLED else None
    yield
    if task:
        task.cancel()


app = FastAPI(title="Job Watch Assistant", lifespan=lifespan)
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"])


@app.get("/api/health")
def health():
    with connect() as c:
        total = c.execute("SELECT COUNT(*) FROM jobs").fetchone()[0]
    return {"status": "ok", "database": DB_PATH, "jobs": total, "scheduler": SCHEDULER_ENABLED, "sources": ["Indeed MCP", "LinkedIn", "Dice", "ZipRecruiter", "jsearch"], "jsearch": {"enabled": JSEARCH_ENABLED, "configured": bool(JSEARCH_RAPIDAPI_KEY), **jsearch_usage()}}


@app.get("/api/dashboard")
def dashboard():
    today = (datetime.now(timezone.utc) - timedelta(days=1)).isoformat()
    with connect() as c:
        total = c.execute("SELECT COUNT(*) FROM jobs").fetchone()[0]
        new_today = c.execute("SELECT COUNT(*) FROM jobs WHERE first_seen>=?", (today,)).fetchone()[0]
        applied = c.execute("SELECT COUNT(*) FROM jobs WHERE status IN ('Applied','Interview','Offer')").fetchone()[0]
        statuses = {r["status"]: r["n"] for r in c.execute("SELECT status, COUNT(*) n FROM jobs GROUP BY status")}
        dist = {r["bucket"]: r["n"] for r in c.execute("SELECT CASE WHEN match_score>=80 THEN '80-100' WHEN match_score>=60 THEN '60-79' WHEN match_score>=40 THEN '40-59' ELSE '0-39' END bucket, COUNT(*) n FROM jobs GROUP BY bucket")}
    return {"new_today": new_today, "total": total, "applied": applied, "statuses": statuses, "score_distribution": dist}


@app.get("/api/jobs")
def jobs(status: str = "", q: str = "", min_score: int = 0, remote_type: str = "", source: str = "", sort: str = "newest", date_range: str = "7d"):
    where, args = [], []
    if status:
        where.append("status=?"); args.append(status)
    if q:
        where.append("(title LIKE ? OR company LIKE ? OR location LIKE ?)"); args += [f"%{q}%"] * 3
    if min_score:
        where.append("match_score>=?"); args.append(min_score)
    if remote_type:
        where.append("remote_type=?"); args.append(remote_type)
    if source:
        where.append("source=?"); args.append(source)
    ranges = {"24h": 1, "3d": 3, "7d": 7, "30d": 30}
    if date_range in ranges:
        cutoff = (datetime.now(timezone.utc) - timedelta(days=ranges[date_range])).isoformat()
        where.append("COALESCE(NULLIF(posted_at,''),first_seen,last_seen)>=?"); args.append(cutoff)
    order_map = {
        "newest": "COALESCE(NULLIF(posted_at,''),first_seen,last_seen) DESC, id DESC",
        "oldest": "COALESCE(NULLIF(posted_at,''),first_seen,last_seen) ASC, id ASC",
        "score_desc": "match_score DESC, COALESCE(NULLIF(posted_at,''),first_seen,last_seen) DESC",
        "score_asc": "match_score ASC, COALESCE(NULLIF(posted_at,''),first_seen,last_seen) DESC",
    }
    order = order_map.get(sort, order_map["newest"])
    sql = "SELECT * FROM jobs" + (" WHERE " + " AND ".join(where) if where else "") + f" ORDER BY {order} LIMIT 500"
    with connect() as c:
        return [rowdict(r) for r in c.execute(sql, args).fetchall()]


@app.get("/api/jsearch/quota")
def jsearch_quota():
    return {"enabled": JSEARCH_ENABLED, "configured": bool(JSEARCH_RAPIDAPI_KEY), **jsearch_usage()}


@app.get("/api/settings")
def get_settings():
    return settings_dict()


@app.post("/api/settings")
def set_settings(payload: SettingsIn):
    return save_settings(payload.model_dump())


@app.post("/api/rescore")
def rescore():
    return {"updated": rescore_existing_jobs()}


@app.get("/api/jobs/{job_id}")
def get_job(job_id: int):
    with connect() as c:
        r = c.execute("SELECT * FROM jobs WHERE id=?", (job_id,)).fetchone()
        if not r:
            raise HTTPException(404, "job not found")
        docs = [dict(x) for x in c.execute("SELECT * FROM generated_docs WHERE job_id=? ORDER BY created_at DESC", (job_id,)).fetchall()]
    d = rowdict(r); d["generated_docs"] = docs
    return d


@app.post("/api/jobs")
def add_job(payload: ManualJob):
    saved, _ = upsert_job(payload.model_dump())
    return saved


@app.post("/api/jobs/{job_id}/status")
def set_status(job_id: int, payload: StatusIn):
    if payload.status not in STATUSES:
        raise HTTPException(400, "bad status")
    with connect() as c:
        c.execute("UPDATE jobs SET status=? WHERE id=?", (payload.status, job_id))
    return get_job(job_id)


@app.post("/api/jobs/{job_id}/notes")
def set_notes(job_id: int, payload: NoteIn):
    with connect() as c:
        c.execute("UPDATE jobs SET notes=? WHERE id=?", (payload.notes, job_id))
    return get_job(job_id)


@app.post("/api/scrape")
def run_scrape():
    return scrape_all()


@app.post("/api/jobs/{job_id}/generate")
def generate(job_id: int, payload: GenerateIn):
    j = get_job(job_id)
    kind = "cover_letter" if payload.kind == "cover_letter" else "tailored_resume"
    content = generate_content(j, kind)
    title = ("Cover Letter" if kind == "cover_letter" else "Tailored Resume") + " - " + j["title"]
    with connect() as c:
        c.execute("INSERT INTO generated_docs(job_id,kind,title,content,created_at) VALUES(?,?,?,?,?)", (job_id, kind, title, content, utcnow()))
        doc_id = c.execute("SELECT last_insert_rowid()").fetchone()[0]
    return {"id": doc_id, "job_id": job_id, "kind": kind, "title": title, "content": content}


@app.get("/api/docs/{doc_id}/export/{fmt}")
def export_doc(doc_id: int, fmt: str):
    with connect() as c:
        d = c.execute("SELECT * FROM generated_docs WHERE id=?", (doc_id,)).fetchone()
    if not d:
        raise HTTPException(404, "doc not found")
    title, content = d["title"], d["content"]
    safe = re.sub(r"[^a-zA-Z0-9_-]+", "-", title).strip("-")[:80] or "document"
    if fmt == "docx":
        out = io.BytesIO(); doc = Document(); doc.add_heading(title, 0)
        for para in content.split("\n"):
            doc.add_paragraph(para.strip())
        doc.save(out)
        return Response(out.getvalue(), media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": f'attachment; filename="{safe}.docx"'})
    if fmt == "pdf":
        out = io.BytesIO(); pdf = SimpleDocTemplate(out, pagesize=letter); styles = getSampleStyleSheet(); story = [Paragraph(html.escape(title), styles["Title"]), Spacer(1, 12)]
        for para in content.split("\n"):
            story.append(Paragraph(html.escape(para) or " ", styles["BodyText"])); story.append(Spacer(1, 6))
        pdf.build(story)
        return Response(out.getvalue(), media_type="application/pdf", headers={"Content-Disposition": f'attachment; filename="{safe}.pdf"'})
    raise HTTPException(400, "fmt must be docx or pdf")


frontend = Path("/app/frontend-dist")
if frontend.exists():
    app.mount("/", StaticFiles(directory=str(frontend), html=True), name="frontend")
